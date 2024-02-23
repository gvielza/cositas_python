# Importar las librerías necesarias
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF

# Función para leer un archivo de texto
def leer_archivo_texto(nombre_archivo):
    with open(nombre_archivo, 'r') as archivo:
        contenido = archivo.read()
    return contenido

# Función para procesar el contenido y contar las palabras
def contar_palabras(texto):
    palabras = texto.split()
    return len(palabras)

# Función para escribir en un archivo de Word
def escribir_en_word(contenido):
    documento = docx.Document()
    # Agregar el contenido al documento
    p = documento.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Conteo de Palabras")
    run.bold = True
    run.font.size = Pt(20)

    p = documento.add_paragraph()
    p.add_run(f"El número de palabras en el texto es: {contenido}").bold = True

    # Guardar el documento
    documento.save("resultado_word.docx")

# Función para escribir en un archivo PDF
def escribir_en_pdf(contenido):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Conteo de Palabras", ln=True, align='C')
    pdf.cell(200, 10, txt=f"El número de palabras en el texto es: {contenido}", ln=True, align='C')
    pdf.output("resultado_pdf.pdf")

# Nombre del archivo de texto
archivo_texto = "texto.txt"

# Leer el archivo de texto
contenido_texto = leer_archivo_texto(archivo_texto)

# Procesar el contenido del archivo de texto
numero_palabras = contar_palabras(contenido_texto)

# Escribir los resultados en un archivo de Word y un archivo PDF
escribir_en_word(numero_palabras)
escribir_en_pdf(numero_palabras)

print("Se han generado los archivos 'resultado_word.docx' y 'resultado_pdf.pdf' con el conteo de palabras.")
